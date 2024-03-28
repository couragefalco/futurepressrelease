import os
import openai
from dotenv import load_dotenv
import json
import streamlit as st
from docx import Document
from io import BytesIO
from docx.shared import Pt

load_dotenv()

openai.api_type = os.getenv("OPENAI_API_TYPE")
openai.api_base = os.getenv("OPENAI_API_BASE")
openai.api_version = os.getenv("OPENAI_API_VERSION")
openai.api_key = os.getenv("OPENAI_API_KEY")

prompt = """
    Schreiben Sie eine ausführliche Pressemitteilung für Igus auf der Grundlage der folgenden Ankündigungsnotizen:
    {notes}
    Stellen Sie sicher, dass die Pressemitteilung umfassend ist, die wichtigsten Errungenschaften hervorhebt und für den Leser ansprechend ist.
    
    Schreiben Sie im JSON-Format:
    {
    "title": "",
    "paragraph1": "",
    "paragraph2": "",
    "paragraph3": ""
    }
    """

#Example Output:
    #{
    #"title": "Record year exceeded thanks to 188,000 active customers - igus increases delivery capacity and grows by 32%.",
    #"paragraph1": "Cologne, Germany, 22nd March 2022 - Customers of motion plastics products - from drive cables to gears - continued to receive a very high proportion of their orders quickly in 2021. The 2019 plan to invest in production and supply chain helped meet the surge in demand. This intense scale-up will continue until 2023, also with a view to the Ukraine conflict. The enterprise's goal gained focus: improve what moves in a CO2-neutral manner, with zero plastic waste and become the easiest company to deal with."
    #"paragraph2": "234 million Euro more sales in one year, with almost the same selling prices until the end of the year, and everything produced as well as sourced in-house - we've never had that before, says Frank Blase, CEO. Our colleagues achieved miracles. And we were lucky to realise our investment plans even in the weak year 2020. The turnover in 2021 amounted to 961 million Euros.",
    #"paragraph3": "32% more turnover than in 2020 also means 26% more than in the record year 2019. This year also saw the beginning of the implementation of the plan that is internally referred to No. 1 Catalogue: more than 80,000 items have since been in stock additionally or in higher quantities. In 15 global distribution centres, the rate of catalogue products shipped the same day or within 24 hours increased to at least 25%. That's probably why the sales growth is almost the same across all product lines, Blase said. The online shops also experienced improvements. Online sales increased by 55 percent in 2021. Customers need to be able to decide immediately on the web whether the plastic solution is usable, and then have it delivered quickly. That's part of the ‘easiest company to deal with’ goal”, according to Frank Blase."
    #}

def generate_press_release(notes):
    try:
        response = openai.ChatCompletion.create(
            engine=os.getenv("OPENAI_ENGINE_ID"),
            response_format={ "type": "json_object" },
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": notes}
            ],
            temperature=0.6,
            max_tokens=1200,
            top_p=0.95,
            frequency_penalty=0,
            presence_penalty=0,
            stop=None
        )
        content = json.loads(response.choices[0].message.content)
        print(content)
        return content
    except Exception as e:
        print(f"Error generating press release: {e}")
        return "An error occurred while generating the press release."



def fill_docx_template(content, template_path="template/pressrelease.docx"):
    doc = Document(template_path)
    # Adjusting function to apply direct formatting to the title
    
    for paragraph in doc.paragraphs:
        if "Title" in paragraph.text:
            paragraph.clear()  # Clear the placeholder text
            run = paragraph.add_run(content["title"])
            run.bold = True  # Make the title bold
            run.font.size = Pt(14)  # Set font size to 14
        elif "First" in paragraph.text:
            paragraph.clear()
            paragraph.add_run(content["paragraph1"])
        elif "Second" in paragraph.text:
            paragraph.clear()
            paragraph.add_run(content["paragraph2"])
        elif "Third" in paragraph.text:  # Assuming this is the correct placeholder text
            paragraph.clear()
            paragraph.add_run(content["paragraph3"])

    return doc


def download_docx(document):
    # Save the docx to a BytesIO object
    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io



def main():
    # Check if essential environment variables are set
    required_env_vars = ["OPENAI_API_TYPE", "OPENAI_API_BASE", "OPENAI_API_VERSION", "OPENAI_API_KEY"]
    missing_vars = [var for var in required_env_vars if not os.getenv(var)]
    
    if missing_vars:
        # Display a message in Streamlit and return early
        missing_vars_str = ", ".join(missing_vars)
        st.error(f"Missing environment variables: {missing_vars_str}. Please set these variables and restart the app.")
        return  # Exit the function early to stop further execution

    st.title('Future Press Release for Igus')
    notes = st.text_area("Notizen deines Berichtes:", height=100)
    if st.button('Erstellung eines Entwurfes'):
        with st.spinner('Entwurf wird entwickelt...'):
            press_release_draft = generate_press_release(notes)
            
            if isinstance(press_release_draft, dict):
                # Display the title and paragraphs
                st.markdown("### Entwurf der Pressemitteilung:")
                st.markdown(f"**{press_release_draft['title']}**")  # Display the title as bold
                st.write(press_release_draft['paragraph1'])  # Display paragraph 1
                st.write(press_release_draft['paragraph2'])  # Display paragraph 2
                st.write(press_release_draft['paragraph3'])  # Display paragraph 3

                # Fill the docx template with the generated content
                doc = fill_docx_template(press_release_draft)
                doc_io = download_docx(doc)

                # Provide a download button for the filled docx
                st.download_button(
                    label="Download Pressemitteilung",
                    data=doc_io,
                    file_name="Generated_Press_Release.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                # Display error message
                st.error("Failed to generate press release. Please try again.")


if __name__ == "__main__":
    main()
